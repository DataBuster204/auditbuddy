from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage


def get_llm(api_key: str):
    return ChatOpenAI(
        model="gpt-4o-mini",
        temperature=0.2,
        openai_api_key=api_key,
    )


# ── 1. SUMMARY ───────────────────────────────────────────────────────────

SUMMARY_SYSTEM = """You are AuditBuddy, a senior financial analyst AI.
Your job is to read financial documents and produce a clear, structured, plain-English summary.

Structure your summary exactly as follows:

## Document Overview
Brief description of what this document is (type, period, entity name if found).

## Key Financial Figures
List the most important numbers: revenue, expenses, net profit/loss, assets, liabilities, equity, cash position, key ratios.

## Performance Snapshot
2-4 sentences on the overall financial health and trend visible in the document.

## Notable Observations
3-5 bullet points of the most important things a finance professional should know.

Be precise. Use the actual numbers from the document. Do not invent figures."""


def summarise(document_text: str, api_key: str) -> str:
    llm = get_llm(api_key)
    messages = [
        SystemMessage(content=SUMMARY_SYSTEM),
        HumanMessage(content=f"Analyse and summarise this financial document:\n\n{document_text[:12000]}")
    ]
    return llm.invoke(messages).content


# ── 2. RED FLAGS ─────────────────────────────────────────────────────────

RED_FLAG_SYSTEM = """You are AuditBuddy, a forensic accounting AI specialising in risk identification.
Analyse the financial document for red flags, anomalies, and areas of concern.

Structure your response exactly as follows:

## Red High-Risk Flags
Critical issues that require immediate attention.

## Moderate Concerns
Issues worth investigating further.

## Areas of Strength
Positive indicators in the financials.

## Recommended Actions
3-5 specific, actionable steps the client or auditor should take.

For each flag, cite the specific figure or section from the document that triggered it."""


def detect_red_flags(document_text: str, api_key: str) -> str:
    llm = get_llm(api_key)
    messages = [
        SystemMessage(content=RED_FLAG_SYSTEM),
        HumanMessage(content=f"Identify red flags in this financial document:\n\n{document_text[:12000]}")
    ]
    return llm.invoke(messages).content


# ── 3. Q&A ───────────────────────────────────────────────────────────────

QA_SYSTEM = """You are AuditBuddy, a financial document assistant.
Answer the user's question accurately and concisely, citing specific figures from the document.

Rules:
- Only answer based on what is in the document
- If the answer is not in the document, say so clearly
- Always reference the specific number or section that supports your answer"""


def answer_question(document_text: str, question: str, api_key: str) -> str:
    llm = get_llm(api_key)
    messages = [
        SystemMessage(content=QA_SYSTEM),
        HumanMessage(content=f"Document:\n{document_text[:12000]}\n\nQuestion: {question}")
    ]
    return llm.invoke(messages).content


# ── 4. AUDIT MEMO ────────────────────────────────────────────────────────

MEMO_SYSTEM = """You are AuditBuddy, a professional audit report writer.
Generate a formal Audit Observation Memo based on the financial document provided.

Use this structure:

---
AUDIT OBSERVATION MEMO

To: Management / Board of Directors
From: AuditBuddy AI Analysis System
Date: {date}
Subject: Financial Document Review

---

## 1. Purpose
## 2. Scope of Review
## 3. Key Findings (each with Observation, Risk Level, Implication, Recommendation)
## 4. Summary Opinion
## 5. Next Steps

Write in formal language appropriate for board-level distribution."""


def generate_audit_memo(document_text: str, api_key: str) -> str:
    from datetime import date
    today = date.today().strftime("%B %d, %Y")
    llm = get_llm(api_key)
    messages = [
        SystemMessage(content=MEMO_SYSTEM.format(date=today)),
        HumanMessage(content=f"Generate the audit memo for this financial document:\n\n{document_text[:12000]}")
    ]
    return llm.invoke(messages).content


# ── 5. CUSTOM REPORT ─────────────────────────────────────────────────────

REPORT_SYSTEM = """You are AuditBuddy, a financial analyst AI.
Generate a thorough, professional report that directly addresses the user's request.
Use clear headings, tables where relevant, and cite specific figures from the document.
Format your response in clean Markdown."""


def generate_custom_report(document_text: str, prompt: str, api_key: str) -> str:
    llm = get_llm(api_key)
    messages = [
        SystemMessage(content=REPORT_SYSTEM),
        HumanMessage(content=f"Document:\n{document_text[:12000]}\n\nAnalysis Request: {prompt}")
    ]
    return llm.invoke(messages).content

# ── EXPORT HELPER ─────────────────────────────────────────────────────────
# This converts the AI output into a Word document for download.
# We already have python-docx installed so no extra packages needed.

def to_word_bytes(content: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from io import BytesIO

    doc = Document()

    # Title
    heading = doc.add_heading(title, 0)
    heading.runs[0].font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

    # Add a line break after title
    doc.add_paragraph("")

    # Process content line by line
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    # Footer note
    doc.add_paragraph("")
    footer = doc.add_paragraph("Generated by AuditBuddy AI | olumidedaramola.dev")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()