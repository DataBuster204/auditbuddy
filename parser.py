import io
import pdfplumber
import pandas as pd
from docx import Document


def parse_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
            tables = page.extract_tables()
            for table in tables:
                if table:
                    try:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        text_parts.append(f"\n[TABLE on Page {i + 1}]\n{df.to_string(index=False)}\n")
                    except Exception:
                        for row in table:
                            text_parts.append("  |  ".join([str(c) for c in row if c]))
    return "\n\n".join(text_parts)


def parse_excel(file_bytes: bytes) -> str:
    text_parts = []
    xl = pd.ExcelFile(io.BytesIO(file_bytes))
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name)
        if df.empty:
            continue
        df = df.dropna(how="all").dropna(axis=1, how="all")
        df = df.fillna("")
        text_parts.append(f"--- Sheet: {sheet_name} ---\n{df.to_string(index=False)}")
    return "\n\n".join(text_parts)


def parse_word(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            try:
                df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
                text_parts.append(f"\n[TABLE {i + 1}]\n{df.to_string(index=False)}\n")
            except Exception:
                for row in rows:
                    text_parts.append("  |  ".join(row))
    return "\n\n".join(text_parts)


def extract_text(uploaded_file) -> tuple[str, str]:
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return parse_pdf(file_bytes), "PDF"
    elif name.endswith((".xlsx", ".xls")):
        return parse_excel(file_bytes), "Excel"
    elif name.endswith(".docx"):
        return parse_word(file_bytes), "Word"
    else:
        raise ValueError(f"Unsupported file type: {uploaded_file.name}")